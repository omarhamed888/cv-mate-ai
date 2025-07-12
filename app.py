from flask import Flask, render_template, request, jsonify, flash, redirect, url_for, session, send_file
import google.generativeai as genai
import os
import PyPDF2
import docx
from werkzeug.utils import secure_filename
import json
import re
import logging
from io import BytesIO
from weasyprint import HTML
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import RGBColor

# Set up logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.secret_key = os.getenv('FLASK_SECRET_KEY', 'your-secret-key-change-this')
app.config['SESSION_TYPE'] = 'filesystem'

# Configure Gemini API
genai.configure(api_key=GEMINI_API_KEY)

# Configure upload settings
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Create upload directory
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(file_path):
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted
    except Exception as e:
        logger.error(f"Error reading PDF: {e}")
    return text

def extract_text_from_docx(file_path):
    text = ""
    try:
        doc = docx.Document(file_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
    except Exception as e:
        logger.error(f"Error reading DOCX: {e}")
    return text

def extract_text_from_txt(file_path):
    text = ""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
    except Exception as e:
        logger.error(f"Error reading TXT: {e}")
    return text

def extract_text_from_file(file_path):
    extension = file_path.rsplit('.', 1)[1].lower()
    if extension == 'pdf':
        return extract_text_from_pdf(file_path)
    elif extension == 'docx':
        return extract_text_from_docx(file_path)
    elif extension == 'txt':
        return extract_text_from_txt(file_path)
    return ""

def extract_contact_info(cv_text):
    """Extract name, email, phone, and LinkedIn from CV text."""
    name = email = phone = linkedin = ""
    lines = cv_text.split('\n')
    
    name_pattern = re.compile(r'^[A-Z][a-z]+(?:\s[A-Z][a-z]+)+', re.MULTILINE)
    email_pattern = re.compile(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}')
    phone_pattern = re.compile(r'\b(\+\d{1,3}\s?)?(\d{3}[-.\s]?){2}\d{4}\b')
    linkedin_pattern = re.compile(r'(?:https?:\/\/)?(?:www\.)?linkedin\.com\/in\/[a-zA-Z0-9-]+')
    
    for line in lines[:50]:  # Limit to first 50 lines to optimize
        line = line.strip()
        if not name and name_pattern.search(line):
            name = name_pattern.search(line).group()
        if not email and email_pattern.search(line):
            email = email_pattern.search(line).group()
        if not phone and phone_pattern.search(line):
            phone = phone_pattern.search(line).group()
        if not linkedin and linkedin_pattern.search(line):
            linkedin = linkedin_pattern.search(line).group()
    
    return name or "[Your Name]", email or "[Your Email]", phone or "[Your Phone]", linkedin or ""

def clean_api_response(response_text):
    cleaned_text = re.sub(r'^```json\n|\n```$', '', response_text, flags=re.MULTILINE)
    return cleaned_text.strip()

def get_generative_model():
    try:
        return genai.GenerativeModel('gemini-1.5-flash')
    except Exception as e:
        logger.error(f"Error initializing Gemini model: {e}")
        return None

model = get_generative_model()

def analyze_cv_with_gemini(cv_text, analysis_type="review", job_description=""):
    if not model:
        logger.error("No valid Gemini model available")
        return {"error": "No valid Gemini model available"}
    
    try:
        # Extract contact info from original CV
        name, email, phone, linkedin = extract_contact_info(cv_text)
        contact_info = f"{name}\n{email}\n{phone}"
        if linkedin:
            contact_info += f"\n{linkedin}"
        
        if analysis_type == "review":
            prompt = f"""
            Analyze this CV and provide a comprehensive review:
            CV Content: {cv_text}
            Provide:
            1. Overall Score (out of 100)
            2. Strengths (list 5-7 key strengths)
            3. Weaknesses (list 5-7 areas for improvement)
            4. Specific recommendations for improvement
            5. Industry-specific feedback
            Format as JSON:
            {{
                "score": number,
                "strengths": [list of strings],
                "weaknesses": [list of strings],
                "recommendations": [list of strings],
                "industry_feedback": "string"
            }}
            """
        elif analysis_type == "ats":
            prompt = f"""
            Analyze this CV for ATS compatibility:
            CV Content: {cv_text}
            Provide:
            1. ATS Score (out of 100)
            2. Keyword optimization score
            3. Format compatibility issues
            4. Missing keywords/skills
            5. Specific ATS improvements needed
            Format as JSON:
            {{
                "ats_score": number,
                "keyword_score": number,
                "format_issues": [list of strings],
                "missing_keywords": [list of strings],
                "improvements": [list of strings]
            }}
            """
        elif analysis_type == "job_match":
            prompt = f"""
            Compare this CV against the job description:
            CV Content: {cv_text}
            Job Description: {job_description}
            Provide:
            1. Match Score (out of 100)
            2. Matched skills and qualifications
            3. Missing requirements
            4. Recommendations to improve match
            5. Gap analysis
            Format as JSON:
            {{
                "match_score": number,
                "matched_skills": [list of strings],
                "missing_requirements": [list of strings],
                "recommendations": [list of strings],
                "gap_analysis": "string"
            }}
            """
        elif analysis_type == "rewrite":
            prompt = f"""
            Rewrite this CV for ATS compatibility and professional formatting:
            Original CV: {cv_text}
            Contact Information: {contact_info}
            Provide:
            1. Improved CV text with the following sections, each clearly labeled in uppercase:
               - CONTACT INFORMATION (include provided name, email, phone, LinkedIn if available)
               - PROFESSIONAL SUMMARY (concise 3-4 sentence overview, keyword-rich)
               - WORK EXPERIENCE (job title, company, location, dates, 3-5 bullet-point achievements per role, max 15 words each)
               - EDUCATION (degree, institution, location, dates, GPA if above 3.0, relevant coursework)
               - SKILLS (list as comma-separated values in a single line, e.g., Python, SQL, Power BI)
               - CERTIFICATIONS (certification name, issuer, dates)
               - VOLUNTEER EXPERIENCE (role, organization, dates, 2-3 bullet-point contributions)
               - PROJECTS (project name, dates, 2-3 bullet-point contributions, tools used)
            2. Summary of changes made
            3. Formatting suggestions for ATS compatibility
            Ensure:
            - Use plain text, no special characters (e.g., replace '&' with 'AND').
            - List skills as a single line with commas (e.g., "Python, SQL, Power BI").
            - Keep bullet points concise (max 15 words) and keyword-rich for ATS.
            - Include projects as a separate section with clear titles, dates, and tools.
            - Avoid tables, headers, footers, or graphics.
            - Reorganize all content from the original CV into the specified sections.
            Format as JSON:
            {{
                "improved_cv": "string",
                "changes_made": [list of strings],
                "formatting_suggestions": [list of strings]
            }}
            """
        response = model.generate_content(prompt)
        cleaned_response = clean_api_response(response.text)
        logger.debug(f"Gemini response for {analysis_type}: {cleaned_response}")
        result = json.loads(cleaned_response)
        if analysis_type == "rewrite" and "improved_cv" not in result:
            logger.error("Gemini response missing 'improved_cv' field")
            return {"error": "Invalid response from Gemini: missing 'improved_cv'"}
        return result
    except Exception as e:
        logger.error(f"Error in analyze_cv_with_gemini: {e}")
        return {"error": str(e)}

def parse_cv_sections(text_content, original_cv_text=""):
    # Define ATS-friendly section headers
    section_headers = [
        "CONTACT INFORMATION",
        "PROFESSIONAL SUMMARY",
        "WORK EXPERIENCE",
        "EDUCATION",
        "SKILLS",
        "CERTIFICATIONS",
        "VOLUNTEER EXPERIENCE",
        "PROJECTS"
    ]
    sections = {header: [] for header in section_headers}
    current_section = None
    fallback_content = []
    
    # Normalize headers to handle variations
    header_map = {
        "SUMMARY": "PROFESSIONAL SUMMARY",
        "EXPERIENCE": "WORK EXPERIENCE",
        "VOLUNTEERING & LEADERSHIP": "VOLUNTEER EXPERIENCE",
        "VOLUNTEER": "VOLUNTEER EXPERIENCE"
    }
    
    lines = text_content.split('\n')
    for i in range(min(200, len(lines))):  # Limit to 200 lines for performance
        line = lines[i].strip()
        if not line:
            continue
        upper_line = line.upper()
        matched_header = upper_line if upper_line in section_headers else header_map.get(upper_line)
        if matched_header:
            current_section = matched_header
        elif current_section:
            sections[current_section].append(line)
        else:
            fallback_content.append(line)
    
    # Fallback: place unparsed content in PROFESSIONAL SUMMARY
    if fallback_content and all(len(content) == 0 for content in sections.values()):
        sections["PROFESSIONAL SUMMARY"] = fallback_content
    
    # Ensure CONTACT INFORMATION is populated
    if not sections["CONTACT INFORMATION"] and original_cv_text:
        name, email, phone, linkedin = extract_contact_info(original_cv_text)
        sections["CONTACT INFORMATION"] = [name, email, phone]
        if linkedin:
            sections["CONTACT INFORMATION"].append(linkedin)
    
    # Handle SKILLS as a single line if comma-separated
    if sections["SKILLS"]:
        skills_line = ", ".join([s.strip() for s in sections["SKILLS"] if s.strip()])
        sections["SKILLS"] = [skills_line] if skills_line else sections["SKILLS"]
    
    # Validate sections
    missing_sections = [header for header, content in sections.items() if not content and header != "VOLUNTEER EXPERIENCE"]
    if missing_sections:
        logger.warning(f"Missing content for sections: {missing_sections}")
    
    logger.debug(f"Parsed sections: {sections}")
    return sections

def generate_ats_pdf(text_content: str, original_cv_text: str = "", margin: int = 20, multi_column_skills: bool = False) -> BytesIO:
    sections = parse_cv_sections(text_content, original_cv_text)
    
    html = f"""
    <html>
    <head>
        <style>
            body {{ font-family: Arial, sans-serif; font-size: 11pt; color: #000000; margin: {margin}px; line-height: 1.3; }}
            h1 {{ font-size: 14pt; font-weight: bold; text-transform: uppercase; margin: 15px 0 8px 0; border-bottom: 1px solid #000000; padding-bottom: 3px; }}
            p {{ margin: 3px 0; }}
            ul {{ list-style-type: disc; margin: 8px 0 8px 25px; }}
            li {{ margin-bottom: 3px; }}
            .section {{ margin-bottom: 20px; }}
            .contact-info {{ text-align: center; margin-bottom: 20px; }}
            .contact-info p {{ margin: 2px 0; }}
            .skills {{ display: {'inline-block' if multi_column_skills else 'block'}; }}
            .skill-item {{ margin-right: 15px; display: inline-block; }}
        </style>
    </head>
    <body>
    """
    
    for section, content in sections.items():
        if content:
            if section == "CONTACT INFORMATION":
                html += '<div class="section contact-info">'
            else:
                html += '<div class="section">'
            html += f'<h1>{section}</h1>'
            if section == "SKILLS" and multi_column_skills:
                html += '<div class="skills">'
                skills = content[0].split(", ")
                for skill in skills:
                    html += f'<span class="skill-item">{skill.strip()}</span>'
                html += '</div>'
            elif section == "SKILLS":
                html += '<div class="skills">'
                for line in content:
                    html += f'<p>{line}</p>'
                html += '</div>'
            elif section in ["WORK EXPERIENCE", "VOLUNTEER EXPERIENCE", "PROJECTS"]:
                html += '<ul>'
                for line in content:
                    html += f'<li>{line}</li>'
                html += '</ul>'
            else:
                for line in content:
                    html += f'<p>{line}</p>'
            html += '</div>'
    
    html += """
    </body>
    </html>
    """
    
    pdf = HTML(string=html).write_pdf()
    return BytesIO(pdf)

def generate_ats_docx(text_content: str, original_cv_text: str = "", margin: int = 0.5, multi_column_skills: bool = False) -> BytesIO:
    doc = Document()
    
    # Set default document style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Define heading style
    try:
        heading_style = doc.styles['Heading 1']
    except KeyError:
        heading_style = doc.styles.add_style('Heading 1', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    heading_style.font.name = 'Arial'
    heading_style.font.size = Pt(14)
    heading_style.font.bold = True
    heading_style.paragraph_format.space_before = Pt(10)
    heading_style.paragraph_format.space_after = Pt(4)
    heading_style.element.get_or_add_pPr().get_or_add_pBdr().append(
        docx.oxml.parse_xml(
            r'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="000000"/>'
        )
    )

    sections = parse_cv_sections(text_content, original_cv_text)
    
    # Set document margins (in inches)
    sections = parse_cv_sections(text_content, original_cv_text)
    doc.sections[0].left_margin = Inches(margin)
    doc.sections[0].right_margin = Inches(margin)
    doc.sections[0].top_margin = Inches(margin)
    doc.sections[0].bottom_margin = Inches(margin)
    
    for section, content in sections.items():
        if content:
            heading = doc.add_heading(section, level=1)
            heading.style = 'Heading 1'
            if section == "CONTACT INFORMATION":
                for line in content:
                    p = doc.add_paragraph(line)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
            elif section == "SKILLS" and multi_column_skills:
                p = doc.add_paragraph()
                run = p.add_run()
                skills = content[0].split(", ")
                for i, skill in enumerate(skills):
                    run.add_text(skill.strip())
                    if i < len(skills) - 1:
                        run.add_text(", ")
                    run.add_break() if (i + 1) % 3 == 0 else None  # 3 columns
                p.paragraph_format.space_after = Pt(4)
            elif section == "SKILLS":
                p = doc.add_paragraph(content[0])
                p.paragraph_format.space_after = Pt(4)
            elif section in ["WORK EXPERIENCE", "VOLUNTEER EXPERIENCE", "PROJECTS"]:
                for line in content:
                    p = doc.add_paragraph(line, style='List Bullet')
                    p.paragraph_format.space_after = Pt(3)
            else:
                for line in content:
                    p = doc.add_paragraph(line)
                    p.paragraph_format.space_after = Pt(3)
            doc.add_paragraph()  # Add spacing between sections
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

@app.route('/download-ats-pdf')
def download_ats_pdf():
    text = session.get('improved_cv', '')
    original_cv = session.get('cv_text', '')
    margin = int(request.args.get('margin', 20))  # Default to 20px
    multi_column = bool(request.args.get('multi_column', False))  # Default to False
    if not text.strip():
        logger.error("No improved CV found in session")
        return "No CV available", 400
    logger.debug(f"Generating PDF with improved_cv: {text[:1000]}..., margin: {margin}, multi_column: {multi_column}")
    pdf_buffer = generate_ats_pdf(text, original_cv, margin, multi_column)
    return send_file(pdf_buffer, as_attachment=True, download_name='ats_friendly_resume.pdf', mimetype='application/pdf')

@app.route('/download-ats-docx')
def download_ats_docx():
    text = session.get('improved_cv', '')
    original_cv = session.get('cv_text', '')
    margin = float(request.args.get('margin', 0.5))  # Default to 0.5 inches
    multi_column = bool(request.args.get('multi_column', False))  # Default to False
    if not text.strip():
        logger.error("No improved CV found in session")
        return "No CV available", 400
    logger.debug(f"Generating DOCX with improved_cv: {text[:1000]}..., margin: {margin}, multi_column: {multi_column}")
    docx_buffer = generate_ats_docx(text, original_cv, margin, multi_column)
    return send_file(docx_buffer, as_attachment=True, download_name='ats_friendly_resume.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/preview-cv', methods=['GET'])
def preview_cv():
    text = session.get('improved_cv', '')
    original_cv = session.get('cv_text', '')
    margin = int(request.args.get('margin', 20))
    multi_column = bool(request.args.get('multi_column', False))
    if not text.strip():
        logger.error("No improved CV found in session for preview")
        return jsonify({'error': 'No CV available'}), 400
    pdf_buffer = generate_ats_pdf(text, original_cv, margin, multi_column)
    pdf_buffer.seek(0)
    return send_file(pdf_buffer, mimetype='application/pdf')

def generate_email_with_gemini(email_type, context):
    try:
        if email_type == "application":
            prompt = f"""
            Generate a professional job application email:
            Context: {context}
            Include:
            1. Appropriate subject line
            2. Professional greeting
            3. Brief introduction
            4. Relevant experience highlights
            5. Professional closing
            Format as JSON:
            {{
                "subject": "string",
                "email_body": "string",
                "tips": [list of strings]
            }}
            """
        elif email_type == "follow_up":
            prompt = f"""
            Generate a professional follow-up email:
            Context: {context}
            Include:
            1. Appropriate subject line
            2. Professional greeting
            3. Polite reminder of previous application
            4. Reaffirm interest
            5. Professional closing
            Format as JSON:
            {{
                "subject": "string",
                "email_body": "string",
                "tips": [list of strings]
            }}
            """
        elif email_type == "thank_you":
            prompt = f"""
            Generate a professional thank you email after interview:
            Context: {context}
            Include:
            1. Appropriate subject line
            2. Professional greeting
            3. Thank you for the interview
            4. Reiterate interest and key points
            5. Professional closing
            Format as JSON:
            {{
                "subject": "string",
                "email_body": "string",
                "tips": [list of strings]
            }}
            """
        response = model.generate_content(prompt)
        cleaned_response = clean_api_response(response.text)
        logger.debug(f"Email generation response: {cleaned_response}")
        return json.loads(cleaned_response)
    except Exception as e:
        logger.error(f"Error in generate_email_with_gemini: {e}")
        return {"error": str(e)}

def generate_interview_questions(cv_text, job_description=""):
    try:
        prompt = f"""
        Generate job-specific interview questions based on this CV and optional job description:
        CV Content: {cv_text}
        Job Description: {job_description}
        Provide:
        1. General questions (3-5)
        2. Technical questions (3-5)
        3. Behavioral questions (3-5)
        Format as JSON:
        {{
            "general_questions": [list of strings],
            "technical_questions": [list of strings],
            "behavioral_questions": [list of strings]
        }}
        """
        response = model.generate_content(prompt)
        cleaned_response = clean_api_response(response.text)
        logger.debug(f"Interview questions response: {cleaned_response}")
        return json.loads(cleaned_response)
    except Exception as e:
        logger.error(f"Error in generate_interview_questions: {e}")
        return {"error": str(e)}

def generate_answer_templates(cv_text, question):
    try:
        prompt = f"""
        Generate answer templates for the following interview question based on this CV:
        CV Content: {cv_text}
        Question: {question}
        Provide 2-3 example answers in professional tone.
        Format as JSON:
        {{
            "question": "string",
            "answers": [list of strings]
        }}
        """
        response = model.generate_content(prompt)
        cleaned_response = clean_api_response(response.text)
        logger.debug(f"Answer templates response: {cleaned_response}")
        return json.loads(cleaned_response)
    except Exception as e:
        logger.error(f"Error in generate_answer_templates: {e}")
        return {"error": str(e)}

@app.route('/download-improved-cv', methods=['GET'])
def download_improved_cv():
    improved_cv = session.get('improved_cv', '')
    if not improved_cv.strip():
        logger.error("No improved CV found in session for download")
        return jsonify({'error': 'No improved CV found'}), 400

    buffer = BytesIO()
    buffer.write(improved_cv.encode('utf-8'))
    buffer.seek(0)
    return send_file(
        buffer,
        as_attachment=True,
        download_name='Improved_CV.txt',
        mimetype='text/plain'
    )

def recommend_career_paths(cv_text):
    try:
        prompt = f"""
        Analyze this CV and recommend career paths:
        CV Content: {cv_text}
        Provide:
        1. Top 3 career paths
        2. Required skills for each path
        3. Steps to transition
        Format as JSON:
        {{
            "career_paths": [
                {{
                    "path": "string",
                    "required_skills": [list of strings],
                    "transition_steps": [list of strings]
                }}
            ]
        }}
        """
        response = model.generate_content(prompt)
        cleaned_response = clean_api_response(response.text)
        logger.debug(f"Career paths response: {cleaned_response}")
        return json.loads(cleaned_response)
    except Exception as e:
        logger.error(f"Error in recommend_career_paths: {e}")
        return {"error": str(e)}

def suggest_projects(cv_text):
    try:
        prompt = f"""
        Suggest small projects to strengthen this CV:
        CV Content: {cv_text}
        Provide:
        1. Project ideas (3-5)
        2. Skills developed
        3. Estimated completion time
        Format as JSON:
        {{
            "projects": [
                {{
                    "idea": "string",
                    "skills_developed": [list of strings],
                    "estimated_time": "string"
                }}
            ]
        }}
        """
        response = model.generate_content(prompt)
        cleaned_response = clean_api_response(response.text)
        logger.debug(f"Project suggestions response: {cleaned_response}")
        return json.loads(cleaned_response)
    except Exception as e:
        logger.error(f"Error in suggest_projects: {e}")
        return {"error": str(e)}

def generate_cv_format(cv_text, format_style):
    try:
        prompt = f"""
        Reformat this CV in the specified style:
        CV Content: {cv_text}
        Format Style: {format_style}
        Provide:
        1. Reformatted CV text
        2. Formatting notes
        Format as JSON:
        {{
            "reformatted_cv": "string",
            "formatting_notes": [list of strings]
        }}
        """
        response = model.generate_content(prompt)
        cleaned_response = clean_api_response(response.text)
        logger.debug(f"CV format response: {cleaned_response}")
        return json.loads(cleaned_response)
    except Exception as e:
        logger.error(f"Error in generate_cv_format: {e}")
        return {"error": str(e)}

def generate_mini_course(cv_text):
    try:
        prompt = f"""
        Generate a mini career course based on this CV:
        CV Content: {cv_text}
        Provide:
        1. Course title
        2. Modules (3-5)
        3. Learning objectives
        Format as JSON:
        {{
            "title": "string",
            "modules": [list of strings],
            "objectives": [list of strings]
        }}
        """
        response = model.generate_content(prompt)
        cleaned_response = clean_api_response(response.text)
        logger.debug(f"Mini course response: {cleaned_response}")
        return json.loads(cleaned_response)
    except Exception as e:
        logger.error(f"Error in generate_mini_course: {e}")
        return {"error": str(e)}

@app.route('/')
def home():
    session['cv_uploaded'] = session.get('cv_uploaded', False)
    session['margin'] = session.get('margin', 20)
    session['multi_column'] = session.get('multi_column', False)
    return render_template('index.html', cv_uploaded=session['cv_uploaded'], margin=session['margin'], multi_column=session['multi_column'])

@app.route('/cv-analysis/review')
def cv_review():
    if not session.get('cv_uploaded', False):
        flash('Please upload a CV to access this feature.', 'warning')
        return redirect(url_for('home'))
    return render_template('cv_review.html')

@app.route('/cv-analysis/ats')
def ats_scanner():
    if not session.get('cv_uploaded', False):
        flash('Please upload a CV to access this feature.', 'warning')
        return redirect(url_for('home'))
    return render_template('ats_scanner.html')

@app.route('/cv-analysis/job-match')
def job_match():
    if not session.get('cv_uploaded', False):
        flash('Please upload a CV to access this feature.', 'warning')
        return redirect(url_for('home'))
    return render_template('job_match.html')

@app.route('/cv-analysis/rewriter')
def cv_rewriter():
    if not session.get('cv_uploaded', False):
        flash('Please upload a CV to access this feature.', 'warning')
        return redirect(url_for('home'))
    return render_template('cv_rewriter.html')

@app.route('/interview-prep/mock-interview')
def mock_interview_page():
    if not session.get('cv_uploaded', False):
        flash('Please upload a CV to access this feature.', 'warning')
        return redirect(url_for('home'))
    return render_template('mock_interview.html')

@app.route('/interview-prep/job-questions')
def job_questions():
    if not session.get('cv_uploaded', False):
        flash('Please upload a CV to access this feature.', 'warning')
        return redirect(url_for('home'))
    return render_template('job_questions.html')

@app.route('/interview-prep/answer-templates')
def answer_templates():
    if not session.get('cv_uploaded', False):
        flash('Please upload a CV to access this feature.', 'warning')
        return redirect(url_for('home'))
    return render_template('answer_templates.html')

@app.route('/upskilling/mini-courses')
def mini_courses():
    if not session.get('cv_uploaded', False):
        flash('Please upload a CV to access this feature.', 'warning')
        return redirect(url_for('home'))
    return render_template('mini_courses.html')

@app.route('/upskilling/career-paths')
def career_paths_page():
    if not session.get('cv_uploaded', False):
        flash('Please upload a CV to access this feature.', 'warning')
        return redirect(url_for('home'))
    return render_template('career_paths.html')

@app.route('/upskilling/project-suggestions')
def project_suggestions_page():
    if not session.get('cv_uploaded', False):
        flash('Please upload a CV to access this feature.', 'warning')
        return redirect(url_for('home'))
    return render_template('project_suggestions.html')

@app.route('/bonus-tools/job-finder')
def job_finder():
    if not session.get('cv_uploaded', False):
        flash('Please upload a CV to access this feature.', 'warning')
        return redirect(url_for('home'))
    return render_template('job_finder.html')

@app.route('/bonus-tools/cv-format')
def cv_format_page():
    if not session.get('cv_uploaded', False):
        flash('Please upload a CV to access this feature.', 'warning')
        return redirect(url_for('home'))
    return render_template('cv_format.html')

@app.route('/bonus-tools/email-writer')
def email_writer():
    if not session.get('cv_uploaded', False):
        flash('Please upload a CV to access this feature.', 'warning')
        return redirect(url_for('home'))
    return render_template('email_writer.html')

@app.route('/upload-cv', methods=['POST'])
def upload_cv():
    if 'cv_file' not in request.files:
        logger.error("No file uploaded in request")
        return jsonify({'error': 'No file uploaded'}), 400
    
    file = request.files['cv_file']
    if file.filename == '':
        logger.error("No file selected")
        return jsonify({'error': 'No file selected'}), 400
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        cv_text = extract_text_from_file(file_path)
        
        try:
            os.remove(file_path)
        except Exception as e:
            logger.error(f"Error removing file {file_path}: {e}")
        
        if not cv_text.strip():
            logger.error("Could not extract text from file")
            return jsonify({'error': 'Could not extract text from file'}), 400
        
        session['cv_text'] = cv_text
        session['cv_uploaded'] = True
        logger.debug(f"CV uploaded successfully, text: {cv_text[:1000]}...")
        return jsonify({'cv_text': cv_text, 'redirect': url_for('home')})
    
    logger.error("Invalid file type")
    return jsonify({'error': 'Invalid file type'}), 400

@app.route('/analyze-cv', methods=['POST'])
def analyze_cv():
    data = request.json
    cv_text = data.get('cv_text', session.get('cv_text', ''))
    analysis_type = data.get('analysis_type', 'review')
    job_description = data.get('job_description', '')
    
    if not cv_text.strip():
        logger.error("CV text is required but empty")
        return jsonify({'error': 'Please upload a CV first'}), 400
    
    logger.debug(f"Analyzing CV with type: {analysis_type}, cv_text: {cv_text[:1000]}...")
    result = analyze_cv_with_gemini(cv_text, analysis_type, job_description)
    if analysis_type == "rewrite" and isinstance(result, dict) and 'improved_cv' in result:
        session['improved_cv'] = result['improved_cv']
        logger.debug(f"Stored improved_cv in session: {result['improved_cv'][:1000]}...")
    else:
        logger.warning(f"Rewrite analysis did not store improved_cv: {result}")
    
    return jsonify(result)

@app.route('/generate-email', methods=['POST'])
def generate_email():
    data = request.json
    email_type = data.get('email_type', 'application')
    context = data.get('context', '')
    
    if not context.strip():
        logger.error("Context is required but empty")
        return jsonify({'error': 'Context is required'}), 400
    
    result = generate_email_with_gemini(email_type, context)
    return jsonify(result)

@app.route('/mock-interview', methods=['POST'])
def mock_interview():
    data = request.json
    cv_text = data.get('cv_text', session.get('cv_text', ''))
    job_description = data.get('job_description', '')
    
    if not cv_text.strip():
        logger.error("CV text is required but empty")
        return jsonify({'error': 'CV text is required'}), 400
    
    result = generate_interview_questions(cv_text, job_description)
    return jsonify(result)

@app.route('/answer-template', methods=['POST'])
def answer_template():
    data = request.json
    cv_text = data.get('cv_text', session.get('cv_text', ''))
    question = data.get('question', '')
    
    if not cv_text.strip() or not question.strip():
        logger.error("CV text or question is missing")
        return jsonify({'error': 'CV text and question are required'}), 400
    
    result = generate_answer_templates(cv_text, question)
    return jsonify(result)

@app.route('/career-paths', methods=['POST'])
def career_paths():
    cv_text = session.get('cv_text', '')
    if not cv_text.strip():
        logger.error("CV text is required but empty")
        return jsonify({'error': 'CV text is required'}), 400
    
    result = recommend_career_paths(cv_text)
    return jsonify(result)

@app.route('/project-suggestions', methods=['POST'])
def project_suggestions():
    cv_text = session.get('cv_text', '')
    if not cv_text.strip():
        logger.error("CV text is required but empty")
        return jsonify({'error': 'CV text is required'}), 400
    
    result = suggest_projects(cv_text)
    return jsonify(result)

@app.route('/cv-format', methods=['POST'])
def cv_format():
    data = request.json
    cv_text = data.get('cv_text', session.get('cv_text', ''))
    format_style = data.get('format_style', 'professional')
    
    if not cv_text.strip():
        logger.error("CV text is required but empty")
        return jsonify({'error': 'CV text is required'}), 400
    
    result = generate_cv_format(cv_text, format_style)
    return jsonify(result)

@app.route('/mini-course', methods=['POST'])
def mini_course():
    cv_text = session.get('cv_text', '')
    if not cv_text.strip():
        logger.error("CV text is required but empty")
        return jsonify({'error': 'CV text is required'}), 400
    
    result = generate_mini_course(cv_text)
    return jsonify(result)

if __name__ == '__main__':
    app.run(debug=True)